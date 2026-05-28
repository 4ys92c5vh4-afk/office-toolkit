#!/usr/bin/env python3
"""
读取 Word 文档 (.docx)
支持：文本提取、元数据读取
"""

import sys
import subprocess
from pathlib import Path


def read_docx_with_pandoc(filepath):
    """使用 pandoc 提取文本（保留格式）"""
    try:
        result = subprocess.run(
            ["pandoc", "--track-changes=all", filepath, "-t", "plain"],
            capture_output=True, text=True, check=True
        )
        return result.stdout
    except FileNotFoundError:
        print("⚠️ pandoc 未安装，尝试用 python-docx")
        return None
    except subprocess.CalledProcessError as e:
        print(f"⚠️ pandoc 错误: {e}")
        return None


def read_docx_with_python(filepath):
    """使用 python-docx 提取文本"""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append(para.text)
        return "\n".join(paragraphs)
    except ImportError:
        print("⚠️ python-docx 未安装，请先运行: pip3 install python-docx")
        return None


def read_docx_metadata(filepath):
    """读取文档元数据"""
    try:
        from docx import Document
        doc = Document(filepath)
        core_props = doc.core_properties
        return {
            "title": core_props.title,
            "author": core_props.author,
            "subject": core_props.subject,
            "created": core_props.created,
            "modified": core_props.modified,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }
    except ImportError:
        return None


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python3 read_docx.py <文件路径>")
        sys.exit(1)

    filepath = sys.argv[1]
    if not Path(filepath).exists():
        print(f"❌ 文件不存在: {filepath}")
        sys.exit(1)

    print(f"📄 读取 Word 文档: {filepath}")
    print("=" * 50)

    # 尝试元数据
    meta = read_docx_metadata(filepath)
    if meta:
        print("\n📋 元数据:")
        for k, v in meta.items():
            print(f"   {k}: {v}")

    # 尝试 pandoc
    text = read_docx_with_pandoc(filepath)
    if text is None:
        text = read_docx_with_python(filepath)

    if text:
        print("\n📝 内容预览 (前 1000 字符):")
        print(text[:1000])
        if len(text) > 1000:
            print(f"\n... (共 {len(text)} 字符)")
