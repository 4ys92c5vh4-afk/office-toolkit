#!/usr/bin/env python3
"""
创建 Word 文档 (.docx)
使用 python-docx 创建带格式、表格、图片的文档
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_simple_docx(output_path):
    """创建简单文档"""
    doc = Document()

    # 标题
    title = doc.add_heading('Office Toolkit 示例文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 段落
    doc.add_paragraph('这是一个使用 python-docx 创建的示例文档。')
    doc.add_paragraph('支持丰富的格式设置，包括：')

    # 列表
    items = ['粗体文本', '斜体文本', '彩色文本', '表格', '图片']
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # 格式化段落
    p = doc.add_paragraph()
    p.add_run('这是粗体 ').bold = True
    p.add_run('这是斜体 ').italic = True
    run = p.add_run('这是红色 ')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.size = Pt(14)

    # 表格
    doc.add_heading('示例表格', level=1)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'

    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '名称'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '大小'

    # 数据行
    data = [
        ['document.docx', 'Word', '25 KB'],
        ['spreadsheet.xlsx', 'Excel', '18 KB'],
    ]
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text

    doc.add_paragraph()
    doc.add_paragraph('--- 文档结束 ---').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    print(f"✅ 文档已创建: {output_path}")


if __name__ == "__main__":
    output = sys.argv[1] if len(sys.argv) > 1 else "output.docx"
    create_simple_docx(output)
